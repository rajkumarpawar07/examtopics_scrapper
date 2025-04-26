import docx
import re
from docx import Document


def extract_questions(doc_path):
    """Extract questions with their numbers from Word document."""
    doc = Document(doc_path)
    questions = []
    current_question_num = None
    current_content = []

    # Pattern to match "Question #: X" where X is a number
    pattern = re.compile(r'Question\s+#:\s+(\d+)')

    for para in doc.paragraphs:
        match = pattern.match(para.text.strip())
        if match:
            # If we already have a question, save it before starting a new one
            if current_question_num is not None:
                questions.append((current_question_num, current_content))
                current_content = []

            # Start a new question
            current_question_num = int(match.group(1))
            current_content.append(para)
        elif current_question_num is not None:
            # Add content to current question
            current_content.append(para)

    # Add the last question if there is one
    if current_question_num is not None and current_content:
        questions.append((current_question_num, current_content))

    return questions


def create_sorted_document(questions, output_path):
    """Create a new document with questions sorted by number."""
    # Sort questions by their number
    sorted_questions = sorted(questions, key=lambda x: x[0])

    # Create a new document
    new_doc = Document()

    for _, paragraphs in sorted_questions:
        for para in paragraphs:
            # Copy paragraph with its formatting
            new_para = new_doc.add_paragraph()
            new_para.text = para.text

            # Copy paragraph formatting
            new_para.style = para.style
            new_para.paragraph_format.alignment = para.paragraph_format.alignment
            new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
            new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing
            new_para.paragraph_format.space_before = para.paragraph_format.space_before
            new_para.paragraph_format.space_after = para.paragraph_format.space_after

            # Copy run formatting (bold, italic, etc.)
            for i, run in enumerate(para.runs):
                if i < len(new_para.runs):
                    new_run = new_para.runs[i]
                else:
                    new_run = new_para.add_run(run.text)

                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
                if run.font.name:
                    new_run.font.name = run.font.name

    # Save the new document
    new_doc.save(output_path)


def sort_questions_in_word_doc(input_path, output_path):
    """Main function to sort questions in a Word document."""
    questions = extract_questions(input_path)
    create_sorted_document(questions, output_path)
    print(f"Sorted document saved to {output_path}")
    print(f"Sorted {len(questions)} questions in ascending order.")


# Usage
if __name__ == "__main__":
    input_file = "./Professional Machine Learning Engineer.docx"  # Change to your input file path
    output_file = "sorted_questions.docx"  # Change to your desired output file path
    sort_questions_in_word_doc(input_file, output_file)