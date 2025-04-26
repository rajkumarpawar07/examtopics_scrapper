from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import json
import time

# Setup Microsoft Edge
options = EdgeOptions()
options.add_argument("--disable-blink-features=AutomationControlled")
service = EdgeService("./msedgedriver.exe")  # <-- Update this path if needed
driver = webdriver.Edge(service=service, options=options)

# Word document
doc = Document()
doc.add_heading('Professional Cloud Database Engineer - Actual Exam Questions', 0)

# JSON list to store all questions
all_questions = []

# Loop through questions
for i in range(1, 132):  # <-- Change range if you want more questions
    query = f"professional database engineer examtopics question {i}"
    driver.get("https://www.google.com/")

    # Search
    search_box = driver.find_element(By.NAME, "q")
    search_box.send_keys(query)
    search_box.send_keys(Keys.RETURN)

    # Get top 2 search results
    WebDriverWait(driver, 10).until(
        EC.presence_of_all_elements_located((By.CSS_SELECTOR, 'h3'))
    )
    search_results = driver.find_elements(By.CSS_SELECTOR, 'h3')[:2]

    valid_question = False

    for result in search_results:
        try:
            result.click()
            WebDriverWait(driver, 5).until(
                EC.presence_of_element_located((By.CLASS_NAME, 'question-discussion-header'))
            )

            header_text = driver.find_element(By.CLASS_NAME, 'question-discussion-header').text
            if "Professional Cloud Database Engineer" in header_text:
                valid_question = True
                break
            else:
                driver.back()
                time.sleep(1)
                search_results = driver.find_elements(By.CSS_SELECTOR, 'h3')[:2]  # refresh
        except Exception as e:
            print(f"[Q{i}] Skipping a result due to error: {e}")
            driver.back()
            time.sleep(1)

    if not valid_question:
        print(f"[Q{i}] Question not found in top 2 results.")
        doc.add_paragraph(f"{i}", style='Heading 2')
        doc.add_paragraph("Question not found.")
        all_questions.append({
            "text": "Question not found.",
            "examId": "profession-database-engineer",
            "id": "",
            "options": [],
            "explanation": ""
        })
        continue

    # Remove popup if exists
    try:
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.ID, 'notRemoverPopup'))
        )
        driver.execute_script("document.getElementById('notRemoverPopup').style.display='none';")
    except:
        print(f"No popup for Question {i}")

    try:
        question_section = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, 'discussion-header-container'))
        )

        question_number = question_section.find_element(By.CSS_SELECTOR, '.question-discussion-header div').text.split('\n')[0]
        question = question_section.find_element(By.CLASS_NAME, 'question-body').find_element(By.TAG_NAME, 'p').text

        options = question_section.find_elements(By.CSS_SELECTOR, '.question-choices-container li')
        options_text = [opt.text for opt in options]

        # Click "Show Suggested Answer" if exists
        try:
            show_answer_button = question_section.find_element(By.CLASS_NAME, "reveal-solution")
            show_answer_button.click()
            time.sleep(1)
        except:
            print(f"[Q{i}] No 'Show Suggested Answer' button found.")

        try:
            answer_element = question_section.find_element(By.CLASS_NAME, 'correct-answer')
            answer = answer_element.text.strip()
        except:
            answer = "Correct Answer: Not found"

        # Write to Word document
        doc.add_paragraph(f"Question #: {question_number}", style='Heading 2')
        doc.add_paragraph(f"Question: {question}", style='Normal')

        doc.add_paragraph("Options:", style='Normal')
        for opt in options_text:
            doc.add_paragraph(opt, style='List Bullet')

        doc.add_paragraph(f"Answer: {answer}", style='Normal')
        doc.add_page_break()

        # Build options properly
        formatted_options = []
        correct_option_letter = answer.replace("Correct Answer:", "").strip()

        for opt in options_text:
            if ". " in opt:
                id_letter, opt_text = opt.split(". ", 1)
                is_correct = (id_letter.strip() == correct_option_letter)
                formatted_options.append({
                    "id": id_letter.strip(),
                    "isCorrect": is_correct,
                    "text": opt_text.strip()
                })

        # Append to JSON
        all_questions.append({
            "text": question,
            "examId": "profession-database-engineer",
            "id": "",
            "options": formatted_options,
            "explanation": ""
        })

    except Exception as e:
        print(f"Error on Question {i}: {e}")

    time.sleep(1)

# Save Word file
doc.save("Professional_Database_Engineer_Questions.docx")

# Save JSON file
with open("Professional_Database_Engineer_Questions.json", "w", encoding='utf-8') as f:
    json.dump(all_questions, f, indent=4, ensure_ascii=False)

# Close browser
driver.quit()
