from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
import time

# Setup Chrome
options = Options()
options.add_argument("--disable-blink-features=AutomationControlled")
service = Service("C:\\Users\\Rajkumar\\Downloads\\chromedriver-win64\\chromedriver-win64\\chromedriver.exe")  # Update this path
driver = webdriver.Chrome(service=service, options=options)

# Word document
doc = Document()
doc.add_heading('Professional Cloud DevOps Engineer Actual Exam Questions', 0)

for i in range(170, 203):
    query = f"exam professional cloud devops engineer question number {i}"
    driver.get("https://www.google.com/")

    # Search
    search_box = driver.find_element(By.NAME, "q")
    search_box.send_keys(query)
    search_box.send_keys(Keys.RETURN)

    # Click the first search result
    # WebDriverWait(driver, 10).until(
    #     EC.presence_of_element_located((By.CSS_SELECTOR, 'h3'))
    # )
    # first_result = driver.find_element(By.CSS_SELECTOR, 'h3')
    # first_result.click()
    # Get top 2 search results
    WebDriverWait(driver, 10).until(
        EC.presence_of_all_elements_located((By.CSS_SELECTOR, 'h3'))
    )
    search_results = driver.find_elements(By.CSS_SELECTOR, 'h3')[:2]

    valid_question = False

    for result in search_results:
        try:
            result.click()
            WebDriverWait(driver, 5).until(
                EC.presence_of_element_located((By.CLASS_NAME, 'question-discussion-header'))
            )

            # Check if it's the correct exam
            header_text = driver.find_element(By.CLASS_NAME, 'question-discussion-header').text
            if "Professional Cloud DevOps Engineer" in header_text:
                valid_question = True
                break
            else:
                driver.back()
                time.sleep(2)
                search_results = driver.find_elements(By.CSS_SELECTOR, 'h3')[:2]  # refresh list
        except Exception as e:
            print(f"[Q{i}] Skipping a result due to error: {e}")
            driver.back()
            time.sleep(2)

    # If none of the top 2 results match, write "Question not found" and skip the rest
    if not valid_question:
        print(f"[Q{i}] Question not found in top 2 results.")
        doc.add_paragraph(f"Question #: {i}", style='Heading 2')
        doc.add_paragraph("Question not found.")
        continue

    # Wait for the popup and remove it
    try:
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.ID, 'notRemoverPopup'))
        )
        driver.execute_script("document.getElementById('notRemoverPopup').style.display='none';")
    except:
        print(f"No popup for Question {i}")

    # Wait for question section to load
    try:
        question_section = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, 'discussion-header-container'))
        )

        # Extract Question Number
        question_number = \
        question_section.find_element(By.CSS_SELECTOR, '.question-discussion-header div').text.split('\n')[0]

        # Extract Question Text
        question = question_section.find_element(By.CLASS_NAME, 'question-body').find_element(By.TAG_NAME, 'p').text

        # Extract Options
        options = question_section.find_elements(By.CSS_SELECTOR, '.question-choices-container li')
        options_text = [opt.text for opt in options]

        # Click "Show Suggested Answer" if it's present
        try:
            show_answer_button = question_section.find_element(By.CLASS_NAME, "reveal-solution")
            show_answer_button.click()
            time.sleep(1)  # wait for answer to be visible
        except:
            print(f"[Q{i}] No 'Show Suggested Answer' button found.")

        # Extract Answer
        try:
            answer_element = question_section.find_element(By.CLASS_NAME, 'correct-answer')
            answer = answer_element.text.strip()
        except:
            answer = "Not found"


        # Extract Answer
        # answer_element = question_section.find_element(By.CLASS_NAME, 'correct-answer')
        # answer = answer_element.text.strip()

        # Write to Word
        doc.add_paragraph(f"{question_number}", style='Heading 2')
        doc.add_paragraph(f"Q: {question}")
        for opt in options_text:
            doc.add_paragraph(opt, style='List Bullet')
        doc.add_paragraph(f"Solution: {answer}", style='Intense Quote')

    except Exception as e:
        print(f"Error on Question {i}: {e}")

    time.sleep(2)  # small pause between iterations

# Save Word File
doc.save("Cloud_DevOps_Engineer_Questions.docx")

# Close browser
driver.quit()
