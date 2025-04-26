import csv
import time
import json
import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document

# File paths
CSV_PATH = "./professional-machine-learning-engineer.csv"
WORD_DOC_PATH = "Professional Machine Learning Engineer.docx"
JSON_PATH = "Professional Machine Learning Engineer.json"
EDGEDRIVER_PATH = "./msedgedriver.exe"

# Initialize output
json_data = []
doc = Document()
doc.add_heading("Google's Professional Cloud Architect Actual Exam Questions", 0)

# Setup Edge WebDriver
def setup_driver():
    options = EdgeOptions()
    options.add_argument("--disable-blink-features=AutomationControlled")
    service = EdgeService(EDGEDRIVER_PATH)
    return webdriver.Edge(service=service, options=options)

# Read URLs from CSV
def read_urls(csv_path):
    with open(csv_path, newline='', encoding='utf-8') as csvfile:
        reader = csv.reader(csvfile)
        next(reader)  # Skip header
        return [row[1] for row in reader if len(row) > 1 and row[1].startswith("http")]

# Hide popup if present
def dismiss_popup(driver):
    try:
        WebDriverWait(driver, 5).until(
            EC.presence_of_element_located((By.ID, 'notRemoverPopup'))
        )
        driver.execute_script("document.getElementById('notRemoverPopup').style.display='none';")
    except:
        pass  # Popup not present

# Extract data from question page
def extract_question_data(driver, url, index):
    try:
        driver.get(url)
        dismiss_popup(driver)

        question_section = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CLASS_NAME, 'discussion-header-container'))
        )

        question_number = question_section.find_element(By.CSS_SELECTOR, '.question-discussion-header div').text.split('\n')[0]
        question_text = question_section.find_element(By.CLASS_NAME, 'question-body').find_element(By.TAG_NAME, 'p').text
        options_elements = question_section.find_elements(By.CSS_SELECTOR, '.question-choices-container li')
        options_text = [opt.text for opt in options_elements]

        # Click 'Show Suggested Answer' if available
        try:
            show_answer_btn = question_section.find_element(By.CLASS_NAME, "reveal-solution")
            show_answer_btn.click()
            time.sleep(1)
        except:
            pass  # No reveal button

        try:
            answer_text = question_section.find_element(By.CLASS_NAME, 'correct-answer').text.strip()
        except:
            answer_text = "Not found"

        # Extract correct option letters using regex
        correct_option_letters = re.findall(r'[A-Z]', answer_text.split(":")[-1].strip().upper())

        # Add to Word Document
        doc.add_paragraph(question_number, style='Heading 2')
        doc.add_paragraph(f"Q: {question_text}")
        for opt in options_text:
            opt_id = opt[0]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(opt)
            if opt_id in correct_option_letters:
                run.bold = True
        doc.add_paragraph(f"Answer: {' '.join(correct_option_letters)}")

        # Add to JSON
        options_json = []
        for opt in options_text:
            opt_id = opt[0]
            clean_text = opt[3:].strip()
            options_json.append({
                "id": opt_id,
                "isCorrect": opt_id in correct_option_letters,
                "text": clean_text
            })

        question_json = {
            "text": question_text,
            "examId": "professional-machine-learning-engineer",
            "id": "",
            "options": options_json,
            "explanation": ""
        }
        return question_json

    except Exception as e:
        print(f"[{index}] Error scraping URL: {url}\nError: {e}")
        return None

# Main script
def main():
    driver = setup_driver()
    urls = read_urls(CSV_PATH)

    for idx, url in enumerate(urls, start=2):
        print(f"[{idx}] Scraping: {url}")
        data = extract_question_data(driver, url, idx)
        if data:
            json_data.append(data)
        time.sleep(1)

    # Save output files
    doc.save(WORD_DOC_PATH)
    with open(JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(json_data, f, indent=4, ensure_ascii=False)

    driver.quit()
    print("✅ Scraping completed and files saved.")

if __name__ == "__main__":
    main()
